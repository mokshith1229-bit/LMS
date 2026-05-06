import streamlit as st
import pandas as pd
from docx import Document
import numpy as np
from io import BytesIO
import re

# --- Page Config ---
st.set_page_config(
    page_title="Training Feedback Report Generator", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS ---
st.markdown("""
<style>
    .block-container { padding-top: 2rem; max-width: 1200px; }
    .section-header { font-size: 1.15rem; font-weight: 600; margin-top: 1.5rem; margin-bottom: 1rem; border-bottom: 1px solid #E2E8F0; padding-bottom: 0.5rem; }
    .mapping-box { background-color: #F8FAFC; padding: 1rem; border-radius: 8px; border: 1px solid #E2E8F0; margin-bottom: 1rem; }
</style>
""", unsafe_allow_html=True)


# -------- Helpers (Backend Logic) --------

def clean_dataframe(df_raw, header_override):
    df = df_raw.copy()
    df = df.dropna(how='all').dropna(axis=1, how='all')
    
    if header_override != "Auto":
        header_idx = int(header_override)
        if header_idx < len(df):
            raw_headers = df.iloc[header_idx].values
            df = df.iloc[header_idx+1:].reset_index(drop=True)
        else:
            raw_headers = df.columns
    else:
        keywords = ["session", "q1", "question", "rating", "batch", "feedback"]
        header_idx = -1
        for i in range(min(10, len(df))):
            row_values = [str(val).lower() for val in df.iloc[i].values]
            if any(any(k in val for k in keywords) for val in row_values):
                header_idx = i
                break
        if header_idx != -1:
            raw_headers = df.iloc[header_idx].values
            df = df.iloc[header_idx+1:].reset_index(drop=True)
        else:
            raw_headers = df.columns
            
    new_headers = []
    seen = {}
    for j, h in enumerate(raw_headers):
        h_str = str(h).strip().lower()
        if pd.isna(h) or h_str in ["", "nan", "none"]:
            h_clean = f"column_{j}"
        else:
            h_clean = h_str
        if h_clean in seen:
            seen[h_clean] += 1
            h_clean = f"{h_clean}_{seen[h_clean]}"
        else:
            seen[h_clean] = 0
        new_headers.append(h_clean)
        
    df.columns = new_headers
    if len(df.columns) > 0:
        df = df.dropna(subset=[df.columns[0]])
    return df

def auto_detect_columns(headers):
    cols = {}
    headers_lower = [str(h).lower() for h in headers]
    for i, h in enumerate(headers_lower):
        if any(k in h for k in ["session", "batch", "training", "class"]):
            if not any(k in h for k in ["comment", "remark", "feedback", "like"]):
                if "session" not in cols: cols["session"] = headers[i]
        for q_idx in range(1, 8):
            q_key = f"q{q_idx}"
            patterns = [rf"\bq\s*{q_idx}\b", rf"question\s*{q_idx}\b", rf"rating\s*{q_idx}\b", rf"^{q_idx}\b"]
            if any(re.search(p, h) for p in patterns) and q_key not in cols:
                cols[q_key] = headers[i]
        if any(k in h for k in ["like", "enjoy", "good", "strength", "positive"]):
            if "likes" not in cols: cols["likes"] = headers[i]
        if any(k in h for k in ["remark", "comment", "improve", "suggestion", "negative"]):
            if "remarks" not in cols: cols["remarks"] = headers[i]
    return cols

def calculate_true_overall_average(df, mapping, layout, num_sessions=None):
    all_numeric_vals = []
    if layout == "Long (Single Session Column)":
        for q in ["q1","q2","q3","q4","q5","q6","q7"]:
            col_name = mapping.get(q)
            if col_name and col_name != "None":
                series = pd.to_numeric(df[col_name], errors='coerce').dropna()
                all_numeric_vals.extend(series.tolist())
    else:
        for s in range(1, num_sessions + 1):
            s_map = mapping.get(f"session_{s}", {})
            for q in ["q1","q2","q3","q4","q5","q6","q7"]:
                col_name = s_map.get(q)
                if col_name and col_name != "None":
                    series = pd.to_numeric(df[col_name], errors='coerce').dropna()
                    all_numeric_vals.extend(series.tolist())
    if all_numeric_vals:
        return round(np.mean(all_numeric_vals), 2)
    return 0.0

def process_data(df, mapping, layout, num_sessions):
    results = {}
    feedback_data = {}
    
    if layout == "Long (Single Session Column)":
        session_col = mapping.get("session")
        if session_col == "None" or not session_col:
            raise ValueError("Session column is required for Long layout.")
            
        sessions = df[session_col].unique()
        for idx, session_val in enumerate(sessions):
            if pd.isna(session_val): continue
            
            session_name = f"Session-{idx+1}"
            sdf = df[df[session_col] == session_val]
            
            q_avgs = []
            session_numeric_vals = []
            for q in ["q1","q2","q3","q4","q5","q6","q7"]:
                col = mapping.get(q)
                if col and col != "None":
                    series = pd.to_numeric(sdf[col], errors='coerce').dropna()
                    if len(series) > 0:
                        q_avgs.append(round(series.mean(), 2))
                        session_numeric_vals.extend(series.tolist())
                    else:
                        q_avgs.append("-")
                else:
                    q_avgs.append("-")
                    
            s_avg = round(np.mean(session_numeric_vals), 2) if session_numeric_vals else 0.0
            results[session_name] = {"count": len(sdf), "q": q_avgs, "avg": s_avg}
            
            likes_col = mapping.get("likes")
            remarks_col = mapping.get("remarks")
            likes = sdf[likes_col].dropna().astype(str).tolist() if likes_col and likes_col != "None" else []
            remarks = sdf[remarks_col].dropna().astype(str).tolist() if remarks_col and remarks_col != "None" else []
            
            feedback_data[session_name] = {
                "likes": [l for l in likes if len(l.strip()) > 2],
                "remarks": [r for r in remarks if len(r.strip()) > 2]
            }
            
    else: # Wide format
        for s in range(1, num_sessions + 1):
            session_name = f"Session-{s}"
            s_map = mapping.get(f"session_{s}", {})
            
            q_avgs = []
            session_numeric_vals = []
            counts = []
            for q in ["q1","q2","q3","q4","q5","q6","q7"]:
                col = s_map.get(q)
                if col and col != "None":
                    series = pd.to_numeric(df[col], errors='coerce').dropna()
                    counts.append(len(series))
                    if len(series) > 0:
                        q_avgs.append(round(series.mean(), 2))
                        session_numeric_vals.extend(series.tolist())
                    else:
                        q_avgs.append("-")
                else:
                    q_avgs.append("-")
                    
            s_avg = round(np.mean(session_numeric_vals), 2) if session_numeric_vals else 0.0
            response_count = max(counts) if counts else 0
            results[session_name] = {"count": response_count, "q": q_avgs, "avg": s_avg}
            
            likes_col = s_map.get("likes")
            remarks_col = s_map.get("remarks")
            likes = df[likes_col].dropna().astype(str).tolist() if likes_col and likes_col != "None" else []
            remarks = df[remarks_col].dropna().astype(str).tolist() if remarks_col and remarks_col != "None" else []
            
            feedback_data[session_name] = {
                "likes": [l for l in likes if len(l.strip()) > 2],
                "remarks": [r for r in remarks if len(r.strip()) > 2]
            }

    overall_avg = calculate_true_overall_average(df, mapping, layout, num_sessions)
    return results, feedback_data, overall_avg

def find_rating_table(doc):
    for table in doc.tables:
        try:
            if not table.rows:
                continue
            first_row_text = " ".join([cell.text.lower() for cell in table.rows[0].cells])
            if "q1" in first_row_text and "average" in first_row_text:
                return table
        except:
            continue
    return None

def fill_row(row, data):
    from docx.oxml import OxmlElement
    while len(row._tr.tc_lst) < 10:
        tc = OxmlElement('w:tc')
        p = OxmlElement('w:p')
        tc.append(p)
        row._tr.append(tc)
    
    cells = row.cells
    cells[1].text = str(data["count"])
    for i, val in enumerate(data["q"]):
        if 2 + i < 9:
            cells[2 + i].text = str(val)
    cells[9].text = str(data["avg"])

def fill_word(doc, results, overall_avg, feedback, extra_fields, limit):
    table = find_rating_table(doc)
    if table is None:
        raise Exception("Ratings table not found! The template must contain a table with 'Q1' and 'Average' in the header.")

    # Match rows exactly by text
    for row in table.rows:
        try:
            cell_text = row.cells[0].text.strip().lower()
            
            for session_name, data in results.items():
                if session_name.lower() == cell_text:
                    fill_row(row, data)
                    
            if "overall" == cell_text or "overall average" in cell_text:
                from docx.oxml import OxmlElement
                while len(row._tr.tc_lst) < 10:
                    tc = OxmlElement('w:tc')
                    p = OxmlElement('w:p')
                    tc.append(p)
                    row._tr.append(tc)
                row.cells[9].text = str(overall_avg)
        except Exception:
            pass

    # Insert Feedback (Strict paragraph appending)
    
    # Pass 1: Likes (under Session-X: labels)
    for session_name, fb in feedback.items():
        session_marker = f"{session_name.lower()}:"
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip().lower()
            
            if text == session_marker:
                p = para._p
                lim = len(fb["likes"]) if limit == "All" else int(limit)
                for idx, val in enumerate(fb["likes"][:lim]):
                    new_p = doc.add_paragraph(f"{idx+1}. {val}")
                    p.addnext(new_p._p)
                    p = new_p._p
                i += lim
            i += 1

    # Pass 2: Remarks (under Remarks: label with session breakdown)
    remarks_marker = "remarks:"
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip().lower()
        
        if text == remarks_marker:
            p = para._p
            for session_name, fb in feedback.items():
                lim = len(fb["remarks"]) if limit == "All" else int(limit)
                if lim > 0 and len(fb["remarks"]) > 0:
                    # Insert Session label
                    new_p_label = doc.add_paragraph(f"{session_name}:")
                    p.addnext(new_p_label._p)
                    p = new_p_label._p
                    
                    # Insert numbered list
                    for idx, val in enumerate(fb["remarks"][:lim]):
                        new_p = doc.add_paragraph(f"{idx+1}. {val}")
                        p.addnext(new_p._p)
                        p = new_p._p
            break # Stop after filling the first Remarks section
        i += 1

    replacements = {
        "[Training Title]": extra_fields["title"],
        "[Batch Number]": extra_fields["batch"],
        "[Date]": extra_fields["date"],
        "[Number of Sessions]": str(len(results))
    }
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if val and key.lower() in para.text.lower():
                para.text = re.sub(re.escape(key), val, para.text, flags=re.IGNORECASE)

    return doc

# -------- Main App UI --------

st.title("Training Feedback Report")
st.markdown('<div class="subtitle">Strict Production Mode: Automated Word document generation from raw training feedback Excel data.</div>', unsafe_allow_html=True)

with st.sidebar:
    st.header("⚙️ Configuration")
    header_override = st.selectbox("Header Row Override", ["Auto"] + [str(i) for i in range(11)])
    data_layout = st.radio("Data Layout", ["Long (Single Session Column)", "Wide (Multiple Session Columns)"])
    
    num_sessions = 0
    if data_layout == "Wide (Multiple Session Columns)":
        num_sessions = st.number_input("Number of Sessions", min_value=1, max_value=20, value=2)
        
    feedback_limit = st.selectbox("Feedback Limit per Session", ["5", "10", "15", "All"], index=1)
    
    st.header("📝 Additional Fields (Optional)")
    training_title = st.text_input("Training Title", "")
    batch_no = st.text_input("Batch Number", "")
    training_date = st.text_input("Date", "")
    extra_fields = {"title": training_title, "batch": batch_no, "date": training_date}

st.markdown('<div class="section-header">1. Upload Files</div>', unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    excel_file = st.file_uploader("Excel/CSV Data", type=["xlsx", "csv"])
with col2:
    word_template = st.file_uploader("Word Template (.docx)", type=["docx"])

if excel_file and word_template:
    try:
        if excel_file.name.endswith("xlsx"):
            df_raw = pd.read_excel(excel_file)
        else:
            df_raw = pd.read_csv(excel_file)

        df = clean_dataframe(df_raw, header_override)
        auto_cols = auto_detect_columns(df.columns)
        all_cols = ["None"] + list(df.columns)
        
        st.markdown('<div class="section-header">2. Required Column Mapping</div>', unsafe_allow_html=True)
        mapping = {}
        def get_index(col_name):
            return all_cols.index(col_name) if col_name in all_cols else 0
            
        with st.container():
            if data_layout == "Long (Single Session Column)":
                st.markdown('<div class="mapping-box">', unsafe_allow_html=True)
                c1, c2, c3 = st.columns(3)
                with c1:
                    mapping["session"] = st.selectbox("Session/Batch Column *", all_cols, index=get_index(auto_cols.get("session")))
                with c2:
                    mapping["likes"] = st.selectbox("Likes/Strengths Column", all_cols, index=get_index(auto_cols.get("likes")))
                with c3:
                    mapping["remarks"] = st.selectbox("Remarks/Improvements Column", all_cols, index=get_index(auto_cols.get("remarks")))
                st.write("**Question Ratings Mapping**")
                q_cols = st.columns(7)
                for i in range(1, 8):
                    with q_cols[i-1]:
                        mapping[f"q{i}"] = st.selectbox(f"Q{i}", all_cols, index=get_index(auto_cols.get(f"q{i}")))
                st.markdown('</div>', unsafe_allow_html=True)
                if mapping["session"] == "None":
                    st.warning("⚠️ You must map the 'Session/Batch Column' to proceed.")
                    st.stop()
            else:
                for s in range(1, num_sessions + 1):
                    st.markdown(f"**Mapping for Session-{s}**")
                    st.markdown('<div class="mapping-box">', unsafe_allow_html=True)
                    s_map = {}
                    c1, c2 = st.columns(2)
                    with c1:
                        s_map["likes"] = st.selectbox(f"Likes Column (Session {s})", all_cols, index=get_index(auto_cols.get("likes")), key=f"l_{s}")
                    with c2:
                        s_map["remarks"] = st.selectbox(f"Remarks Column (Session {s})", all_cols, index=get_index(auto_cols.get("remarks")), key=f"r_{s}")
                    q_cols = st.columns(7)
                    for i in range(1, 8):
                        with q_cols[i-1]:
                            s_map[f"q{i}"] = st.selectbox(f"Q{i}", all_cols, index=0, key=f"q_{s}_{i}")
                    mapping[f"session_{s}"] = s_map
                    st.markdown('</div>', unsafe_allow_html=True)

        results, feedback_data, overall_avg = process_data(df, mapping, data_layout, num_sessions)

        if not results:
            st.error("No valid session data could be extracted based on the current mapping.")
            st.stop()

        st.markdown('<div class="section-header">3. Processing Preview</div>', unsafe_allow_html=True)
        m1, m2, m3 = st.columns(3)
        with m1:
            st.metric("True Overall Average", f"{overall_avg} / 5.0")
        with m2:
            st.metric("Total Sessions", len(results))
        with m3:
            st.metric("Total Responses", sum(v["count"] for v in results.values()))

        st.markdown("**Session-wise Breakdown**")
        preview_list = []
        for session, data in results.items():
            row = {"Session": session, "Count": data["count"], "Avg": data["avg"]}
            for i, val in enumerate(data["q"]):
                row[f"Q{i+1}"] = val
            preview_list.append(row)
        st.dataframe(pd.DataFrame(preview_list), use_container_width=True, hide_index=True)

        st.markdown('<div class="section-header">4. Download Output</div>', unsafe_allow_html=True)
        with st.spinner("Compiling Word document..."):
            doc = Document(word_template)
            doc = fill_word(doc, results, overall_avg, feedback_data, extra_fields, feedback_limit)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.success("Report successfully generated!")
            st.download_button(
                label="⬇️ Download Final Report (.docx)",
                data=buffer,
                file_name="training_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
else:
    st.info("Awaiting both Excel data and Word template to begin processing.")
